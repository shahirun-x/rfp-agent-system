from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
import shutil
import os
import tempfile
from docx import Document as DocxDocument # Renamed to avoid conflict

# --- AI & DB Imports ---
from langchain_groq import ChatGroq
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv

# --- UPGRADE #5 SPECIFIC IMPORTS ---
from llama_parse import LlamaParse
from langchain_core.documents import Document as LangChainDocument # <--- FIXED IMPORT
import nest_asyncio 
nest_asyncio.apply()

load_dotenv()

app = FastAPI(title="RFP Analyst API", version="3.0 (LlamaParse)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- CONFIGURATION ---
PINECONE_INDEX_NAME = "rfp-agent" 

# --- HELPER: Connect to Cloud DB ---
def get_vectorstore():
    embeddings = FastEmbedEmbeddings()
    return PineconeVectorStore.from_existing_index(PINECONE_INDEX_NAME, embeddings)

# --- MODELS ---
class ChatRequest(BaseModel):
    question: str
    history: List[dict] = []

class DownloadRequest(BaseModel):
    history: List[dict] = []

class RevisionRequest(BaseModel):
    original_text: str
    feedback: str

# --- ENDPOINTS ---

@app.get("/")
def read_root():
    return {"status": "Active", "service": "RFP Agent (LlamaParse Vision)"}

@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    # 1. Validation
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed.")
    
    # 2. Save Temp
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        shutil.copyfileobj(file.file, tmp_file)
        tmp_path = tmp_file.name
    
    try:
        print(f"📄 Processing {file.filename} with LlamaParse (Vision AI)...")
        
        # --- NEW: LlamaParse Logic ---
        parser = LlamaParse(
            api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
            result_type="markdown", 
            verbose=True
        )
        
        # LlamaParse is async
        llama_docs = await parser.aload_data(tmp_path)
        
        # CONVERSION STEP: LlamaIndex -> LangChain
        docs = []
        for d in llama_docs:
            docs.append(LangChainDocument(
                page_content=d.text,
                metadata={"source": file.filename}
            ))
            
        print(f"✅ Parsed {len(docs)} pages of rich text.")
        
        # Split text
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(docs)
        
        # Upload to Pinecone
        print("🌲 Uploading to Pinecone...")
        embeddings = FastEmbedEmbeddings()
        
        PineconeVectorStore.from_documents(
            documents=splits, 
            embedding=embeddings, 
            index_name=PINECONE_INDEX_NAME
        )
        
        return {"message": "PDF parsed with LlamaParse and saved to Cloud", "chunks": len(splits)}
    
    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        os.remove(tmp_path)

@app.post("/chat")
async def chat_endpoint(request: ChatRequest):
    try:
        vector_store = get_vectorstore()
    except Exception as e:
        raise HTTPException(status_code=500, detail="Could not connect to Pinecone. Check API Key.")
    
    question = request.question
    history = request.history
    
    chat_history_text = ""
    for msg in history[-4:]:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        chat_history_text += f"{role.upper()}: {content}\n"

    llm_router = ChatGroq(model="llama-3.1-8b-instant", temperature=0)
    router_prompt = f"""
    Classify the user's intent based on Question and History.
    If it mentions cost/risk/business -> LEGAL.
    If it mentions code/architecture -> TECHNICAL.
    Return ONLY the category word.
    
    History: {chat_history_text}
    Question: {question}
    """
    category = llm_router.invoke(router_prompt).content.strip().upper()
    if "LEGAL" in category: category = "LEGAL"
    else: category = "TECHNICAL"
    
    retriever = vector_store.as_retriever()
    docs = retriever.invoke(question)
    
    sources = set()
    for doc in docs:
        # LlamaParse sometimes doesn't give page numbers, so we default to "Source File"
        page_ref = doc.metadata.get("page", "Parsed Data")
        sources.add(f"{page_ref}")
    sources_list = sorted(list(sources))
    
    context = "\n".join([d.page_content for d in docs])
    
    if category == "TECHNICAL":
        system = "You are a Technical Architect. Answer with technical depth."
    else:
        system = "You are a Business Analyst. Answer focusing on risk and value."
        
    final_prompt = f"""
    {system}
    Context: {context}
    History: {chat_history_text}
    Question: {question}
    Answer:
    """
    response = llm_router.invoke(final_prompt)
    
    return {
        "category": category,
        "answer": response.content,
        "sources": sources_list
    }

@app.post("/generate-brief")
async def generate_brief():
    vector_store = get_vectorstore()
    retriever = vector_store.as_retriever(search_kwargs={"k": 10})
    docs = retriever.invoke("Overview executive summary risks technical architecture")
    context = "\n".join([d.page_content for d in docs])
    
    llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.3)
    prompt = f"""
    You are a Senior Proposal Writer. Write a structured Executive Brief.
    Use Markdown headers (#).
    
    Context: {context}
    """
    response = llm.invoke(prompt)
    
    return {"answer": response.content, "category": "WRITER"}

@app.post("/refine-brief")
async def refine_brief(request: RevisionRequest):
    llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.3)
    prompt = f"""
    You are a Senior Editor. Rewrite this draft based on feedback.
    Draft: {request.original_text}
    Feedback: {request.feedback}
    """
    response = llm.invoke(prompt)
    return {"answer": response.content}

@app.post("/download-report")
async def download_report(request: DownloadRequest):
    doc = DocxDocument()
    doc.add_heading('RFP Analysis Report', 0)
    for msg in request.history:
        role = msg.get("role", "user").upper()
        content = msg.get("content", "")
        doc.add_heading(role, level=2)
        doc.add_paragraph(content)
        doc.add_paragraph("---")
    
    temp_filename = "rfp_report.docx"
    doc.save(temp_filename)
    return FileResponse(path=temp_filename, filename="RFP_Analysis.docx")